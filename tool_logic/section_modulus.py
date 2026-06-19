from flask import Blueprint, request, jsonify, Response, send_file
import io
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

section_modulus_bp = Blueprint('section_modulus', __name__)

# ════════════════════════════════════════════════
# CORE GEOMETRICAL CALCULATION ENGINES
# ════════════════════════════════════════════════

def _calculate_t_section(hw, tw, bf, tf):
    a1 = hw * tw
    y1 = hw / 2.0
    a2 = bf * tf
    y2 = hw + (tf / 2.0)
    
    area = a1 + a2
    y_bar = ((a1 * y1) + (a2 * y2)) / area if area > 0 else 0
    
    i1 = (tw * (hw ** 3)) / 12.0
    i2 = (bf * (tf ** 3)) / 12.0
    ix = (i1 + a1 * ((y1 - y_bar) ** 2)) + (i2 + a2 * ((y2 - y_bar) ** 2))
    
    y_bottom = y_bar
    z_bottom = ix / y_bottom if y_bottom > 0 else 0
    return {'area': round(area, 2), 'centroid': round(y_bar, 2), 'moment_of_inertia': round(ix, 2), 'z_bottom': round(z_bottom, 2)}

def _calculate_l_angle(hw, tw, bf, tf):
    a1 = hw * tw
    y1 = hw / 2.0
    a2 = (bf - tw) * tf
    y2 = tf / 2.0
    
    area = a1 + a2
    y_bar = ((a1 * y1) + (a2 * y2)) / area if area > 0 else 0
    
    i1 = (tw * (hw ** 3)) / 12.0
    i2 = ((bf - tw) * (tf ** 3)) / 12.0
    ix = (i1 + a1 * ((y1 - y_bar) ** 2)) + (i2 + a2 * ((y2 - y_bar) ** 2))
    
    y_bottom = y_bar
    z_bottom = ix / y_bottom if y_bottom > 0 else 0
    return {'area': round(area, 2), 'centroid': round(y_bar, 2), 'moment_of_inertia': round(ix, 2), 'z_bottom': round(z_bottom, 2)}

def _calculate_i_beam(hw, tw, bf, tf):
    a_web = hw * tw
    y_web = tf + (hw / 2.0)
    a_flange = bf * tf
    y_bottom_flange = tf / 2.0
    y_top_flange = tf + hw + (tf / 2.0)
    
    area = a_web + (2 * a_flange)
    total_height = hw + (2 * tf)
    y_bar = total_height / 2.0
    
    i_web = (tw * (hw ** 3)) / 12.0
    i_flange = (bf * (tf ** 3)) / 12.0
    
    ix = (i_web + a_web * ((y_web - y_bar) ** 2)) + \
         (i_flange + a_flange * ((y_bottom_flange - y_bar) ** 2)) + \
         (i_flange + a_flange * ((y_top_flange - y_bar) ** 2))
         
    z_bottom = ix / y_bar if y_bar > 0 else 0
    return {'area': round(area, 2), 'centroid': round(y_bar, 2), 'moment_of_inertia': round(ix, 2), 'z_bottom': round(z_bottom, 2)}

def _calculate_c_channel(hw, tw, bf, tf):
    a_web = hw * tw
    y_web = tf + (hw / 2.0)
    a_flange = bf * tf
    y_bot_f = tf / 2.0
    y_top_f = tf + hw + (tf / 2.0)
    
    area = a_web + (2 * a_flange)
    total_height = hw + (2 * tf)
    y_bar = total_height / 2.0
    
    i_web = (tw * (hw ** 3)) / 12.0
    i_flange = (bf * (tf ** 3)) / 12.0
    
    ix = (i_web + a_web * ((y_web - y_bar) ** 2)) + \
         (i_flange + a_flange * ((y_bot_f - y_bar) ** 2)) + \
         (i_flange + a_flange * ((y_top_f - y_bar) ** 2))
         
    z_bottom = ix / y_bar if y_bar > 0 else 0
    return {'area': round(area, 2), 'centroid': round(y_bar, 2), 'moment_of_inertia': round(ix, 2), 'z_bottom': round(z_bottom, 2)}

def _calculate_flat_bar(h, t):
    area = h * t
    y_bar = h / 2.0
    ix = (t * (h ** 3)) / 12.0
    z_bottom = ix / y_bar if y_bar > 0 else 0
    return {'area': round(area, 2), 'centroid': round(y_bar, 2), 'moment_of_inertia': round(ix, 2), 'z_bottom': round(z_bottom, 2)}

def _calculate_bulb_flat(hw, tw, db, tb):
    a1 = hw * tw
    y1 = hw / 2.0
    a2 = db * tb
    y2 = hw + (tb / 2.0)
    
    area = a1 + a2
    y_bar = ((a1 * y1) + (a2 * y2)) / area if area > 0 else 0
    
    i1 = (tw * (hw ** 3)) / 12.0
    i2 = (db * (tb ** 3)) / 12.0
    ix = (i1 + a1 * ((y1 - y_bar) ** 2)) + (i2 + a2 * ((y2 - y_bar) ** 2))
    
    z_bottom = ix / y_bar if y_bar > 0 else 0
    return {'area': round(area, 2), 'centroid': round(y_bar, 2), 'moment_of_inertia': round(ix, 2), 'z_bottom': round(z_bottom, 2)}

# ════════════════════════════════════════════════
# DYNAMIC ROUTE ENDPOINTS
# ════════════════════════════════════════════════

@section_modulus_bp.route('/api/t-section/calculate', methods=['POST'])
def t_section_calculate():
    d = request.get_json(silent=True) or request.form.to_dict()
    try:
        if not d: 
            return jsonify({'ok': False, 'error': "No input data received."}), 400
            
        profile = d.get('profile', 'T')
        
        def get_float(key):
            val = d.get(key)
            if val is None or str(val).strip() == "":
                raise ValueError(f"Missing or empty value for field: {key}")
            return float(val)
        
        if profile == 'T':
            results = _calculate_t_section(get_float('web_height'), get_float('web_thickness'), get_float('flange_width'), get_float('flange_thickness'))
        elif profile == 'L':
            results = _calculate_l_angle(get_float('web_height'), get_float('web_thickness'), get_float('flange_width'), get_float('flange_thickness'))
        elif profile == 'I':
            results = _calculate_i_beam(get_float('web_height'), get_float('web_thickness'), get_float('flange_width'), get_float('flange_thickness'))
        elif profile == 'C':
            results = _calculate_c_channel(get_float('web_height'), get_float('web_thickness'), get_float('flange_width'), get_float('flange_thickness'))
        elif profile == 'FB':
            results = _calculate_flat_bar(get_float('bar_height'), get_float('bar_thickness'))
        elif profile == 'BF':
            if d.get('bulb_size') != 'others':
                bulb_presets = {
                    'hp80':  {'hw': 80,  'tw': 6,  'db': 18, 'tb': 9},
                    'hp100': {'hw': 100, 'tw': 7,  'db': 22, 'tb': 11},
                    'hp120': {'hw': 120, 'tw': 8,  'db': 26, 'tb': 13}
                }
                p = bulb_presets.get(d.get('bulb_size'), bulb_presets['hp80'])
                results = _calculate_bulb_flat(p['hw'], p['tw'], p['db'], p['tb'])
            else:
                results = _calculate_bulb_flat(get_float('web_height'), get_float('web_thickness'), get_float('bulb_diameter'), get_float('bulb_thickness'))
        else:
            return jsonify({'ok': False, 'error': f"Unknown profile type: {profile}"}), 400

        return jsonify({'ok': True, 'results': results})
        
    except (KeyError, ValueError, TypeError) as e:
        return jsonify({'ok': False, 'error': f"Invalid inputs: {str(e)}"}), 400
    except Exception as e:
        return jsonify({'ok': False, 'error': f"Internal Server Error: {str(e)}"}), 500

@section_modulus_bp.route('/api/section-modulus/export-excel', methods=['POST'])
def export_excel():
    try:
        d = request.get_json(silent=True) or request.form.to_dict()
        if not d:
            return jsonify({'ok': False, 'error': 'No data provided'}), 400

        output = io.StringIO()
        output.write("NEXUS MARINE DESIGN SUITE - ENGINEERING REPORT\n")
        output.write(f"Profile Type, {d.get('profile', 'T Section')}\n\n")
        output.write("Property, Value, Unit\n")
        output.write(f"Section Modulus (Bottom), {d.get('z_bottom', '0')}, cm³\n")
        output.write(f"Total Cross-Sectional Area, {d.get('area', '0')}, cm²\n")
        output.write(f"Centroid Position (y-bar), {d.get('centroid', '0')}, mm\n")
        output.write(f"Moment of Inertia (Ix), {d.get('moment_of_inertia', '0')}, cm⁴\n")
        
        output.seek(0)
        return Response(
            output.getvalue(),
            mimetype="text/csv",
            headers={"Content-disposition": "attachment; filename=Nexus_Engineering_Report.csv"}
        )
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

def _calculate_custom_polygon(vertices):
    n = len(vertices)
    if n < 3:
        return {'area': 0, 'centroid': 0, 'moment_of_inertia': 0, 'z_bottom': 0}

    area = 0.0
    cx = 0.0
    cy = 0.0
    ix_raw = 0.0

    for i in range(n):
        x0, y0 = vertices[i]['x'], vertices[i]['y']
        x1, y1 = vertices[(i + 1) % n]['x'], vertices[(i + 1) % n]['y']

        factor = (x0 * y1 - x1 * y0)
        area += factor
        cx += (x0 + x1) * factor
        cy += (y0 + y1) * factor
        ix_raw += (y0**2 + y0*y1 + y1**2) * factor

    area = area / 2.0
    if abs(area) == 0:
        return {'area': 0, 'centroid': 0, 'moment_of_inertia': 0, 'z_bottom': 0}

    cx = cx / (6.0 * area)
    cy = cy / (6.0 * area)

    ix_centroidal = (ix_raw / 12.0) - (area * (cy ** 2))

    area = abs(area)
    ix_centroidal = abs(ix_centroidal)
    z_bottom = ix_centroidal / cy if cy > 0 else 0

    return {
        'area': round(area, 2),
        'centroid': round(cy, 2),
        'moment_of_inertia': round(ix_centroidal, 2),
        'z_bottom': round(z_bottom, 2)
    }

@section_modulus_bp.route('/api/advanced-section/calculate', methods=['POST'])
def advanced_section_calculate():
    d = request.get_json(silent=True) or {}
    vertices = d.get('vertices', [])
    
    if not vertices or len(vertices) < 3:
        return jsonify({'ok': False, 'error': "Please plot at least 3 structural coordinate nodes."}), 400
        
    try:
        results = _calculate_custom_polygon(vertices)
        return jsonify({'ok': True, 'results': results})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

# ════════════════════════════════════════════════
# CORPORATE WORD DOCX REPORT GENERATION ENGINE
# ════════════════════════════════════════════════

@section_modulus_bp.route('/api/section-modulus/export-word', methods=['POST'])
def export_word():
    try:
        d = request.get_json(silent=True) or request.form.to_dict()
        if not d:
            return jsonify({'ok': False, 'error': 'No data provided'}), 400

        profile = d.get('profile', 'T Section')
        z_bottom = d.get('z_bottom', '--')
        area = d.get('area', '--')
        centroid = d.get('centroid', '--')
        moment_of_inertia = d.get('moment_of_inertia', '--')
        
        web_height = d.get('web_height', '--')
        web_thickness = d.get('web_thickness', '--')
        flange_width = d.get('flange_width', '--')
        flange_thickness = d.get('flange_thickness', '--')

        formatted_date = datetime.date.today().strftime('%d-%m-%Y')

        doc = Document()

        # Page Margin Configuration
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        COLOR_BLUE = RGBColor(43, 120, 197)        # CSL Primary Accent Blue
        COLOR_DARK_TEXT = RGBColor(17, 17, 17)     # Titles/Headings
        COLOR_MUTED_TEXT = RGBColor(110, 110, 110) # Metadata tags
        HEX_BORDER_BLUE = "5B9BD5"                 # Table Outer Borders
        HEX_LIGHT_BG = "F2F7FC"                    # Header shading
        COLOR_CHARCOAL = RGBColor(51, 51, 51)

        def set_cell_shading(cell, fill_hex):
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def set_custom_table_borders(table, border_hex):
            tblPr = table._tbl.tblPr
            borders_elm = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_hex}"/>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_hex}"/>'
                f'  <w:left w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
                f'  <w:insideV w:val="none"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders_elm)

        def append_running_page_header(document):
            tbl = document.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            tbl.columns[0].width = Inches(3.5)
            tbl.columns[1].width = Inches(3.5)

            cell_l = tbl.rows[0].cells[0]
            p_l = cell_l.paragraphs[0]
            p_l.paragraph_format.space_after = Pt(2)
            r_l = p_l.add_run("-Drawing Title-\n-Drawing No-\n-Rev. /Date-")
            r_l.font.size = Pt(8.5)
            r_l.font.color.rgb = COLOR_MUTED_TEXT

            cell_r = tbl.rows[0].cells[1]
            p_r = cell_r.paragraphs[0]
            p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_r.paragraph_format.space_after = Pt(2)
            r_r = p_r.add_run("-Project No.-\n-Yard No.-\n-Vessel/Project Name-  ")
            r_r.font.size = Pt(8.5)
            r_r.font.color.rgb = COLOR_MUTED_TEXT
            
            try:
                p_r.add_run().add_picture('static/images/csl_logo.png', width=Inches(0.45))
            except Exception:
                pass

            p_line = document.add_paragraph()
            p_line.paragraph_format.space_before = Pt(0)
            p_line.paragraph_format.space_after = Pt(18)
            p_line_run = p_line.add_run("━" * 65)
            p_line_run.font.color.rgb = COLOR_BLUE
            p_line_run.font.size = Pt(8)

        # ==========================================================
        # PAGE 1: SYSTEM COVER LAYOUT
        # ==========================================================
        p_logo = doc.add_paragraph()
        try:
            p_logo.add_run().add_picture('static/images/csl_logo.png', width=Inches(0.95))
        except Exception:
            pass

        p_corp = doc.add_paragraph()
        r_corp = p_corp.add_run("COCHIN SHIPYARD LTD.\nA GOVT. OF INDIA ENTERPRISE")
        r_corp.font.name = 'Arial'
        r_corp.font.size = Pt(13)
        r_corp.font.bold = True
        r_corp.font.color.rgb = COLOR_BLUE
        
        p_line_top = doc.add_paragraph()
        r_line_top = p_line_top.add_run("━" * 65)
        r_line_top.font.color.rgb = COLOR_BLUE

        p_proj_title = doc.add_paragraph()
        p_proj_title.paragraph_format.space_before = Pt(45)
        p_proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_proj = p_proj_title.add_run("NEXUS: MARINE DESIGN SUITE")
        r_proj.font.name = 'Arial'
        r_proj.font.size = Pt(22)
        r_proj.font.bold = True
        r_proj.font.color.rgb = COLOR_DARK_TEXT

        p_vessel = doc.add_paragraph()
        p_vessel.paragraph_format.space_before = Pt(75)
        p_vessel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ves = p_vessel.add_run("-Insert Vessel Picture-")
        r_ves.font.name = 'Arial'
        r_ves.font.size = Pt(11)
        r_ves.font.italic = True
        r_ves.font.color.rgb = COLOR_MUTED_TEXT

        p_drg_title = doc.add_paragraph()
        p_drg_title.paragraph_format.space_before = Pt(75)
        p_drg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_drg = p_drg_title.add_run("SECTION PROPERTIES CALCULATION REPORT")
        r_drg.font.name = 'Arial'
        r_drg.font.size = Pt(16)
        r_drg.font.bold = True
        r_drg.font.color.rgb = COLOR_DARK_TEXT

        table_cover = doc.add_table(rows=4, cols=4)
        table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_custom_table_borders(table_cover, HEX_BORDER_BLUE)
        table_cover.autofit = False
        table_cover.columns[0].width = Inches(1.5)
        table_cover.columns[1].width = Inches(2.0)
        table_cover.columns[2].width = Inches(1.5)
        table_cover.columns[3].width = Inches(2.0)
        
        meta_matrix = [
            ["Drg. No.", ": CSL-ST-MOD-001", "Department", ": Basic Design"],
            ["Project", f": {profile}", "Document Type", ": Calculation Report"],
            ["Yard No.", ": BY-1002", "Owner", ": Technical Records Section"],
            ["Class", ": IRS / DNV GL Dual Class", "Pages", f": {formatted_date}"]
        ]
        
        for idx, row_vals in enumerate(meta_matrix):
            row = table_cover.rows[idx]
            for c_idx, val_text in enumerate(row_vals):
                cell = row.cells[c_idx]
                cell.text = val_text
                p = cell.paragraphs[0]
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(9.5)
                if c_idx % 2 == 0:
                    p.runs[0].font.bold = True

        p_disc_notice = doc.add_paragraph()
        p_disc_notice.paragraph_format.space_before = Pt(35)
        p_disc_notice.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_notice = p_disc_notice.add_run(
            "This document/specification is the property of Cochin Shipyard Limited, and it must not be copied "
            "or the contents thereof or any information received in conjunction therewith must not be imparted/shared "
            "to any third party or utilized for any other purpose. The receipt of the document/specification implies "
            "that the conditions as mentioned herein are accepted."
        )
        r_notice.font.size = Pt(8.5)
        r_notice.font.italic = True
        r_notice.font.color.rgb = COLOR_CHARCOAL

        # ==========================================================
        # PAGE 2: REVISION CONTROL SHEET BLOCK
        # ==========================================================
        doc.add_page_break()
        append_running_page_header(doc)

        p_rev_head = doc.add_paragraph()
        r_rev_head = p_rev_head.add_run("Revision Control Sheet")
        r_rev_head.font.name = 'Arial'
        r_rev_head.font.size = Pt(14)
        r_rev_head.font.bold = True
        r_rev_head.font.color.rgb = COLOR_DARK_TEXT

        table_rev = doc.add_table(rows=2, cols=7)
        table_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_custom_table_borders(table_rev, "000000")
        
        rev_headers = ["S. No", "Revision", "Date", "Description", "Prepared", "Checked", "Approved"]
        for idx, header_text in enumerate(rev_headers):
            cell = table_rev.rows[0].cells[idx]
            cell.text = header_text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_shading(cell, HEX_LIGHT_BG)

        row_cells = table_rev.rows[1].cells
        row_cells[0].text = "1."
        row_cells[1].text = "0"
        row_cells[2].text = formatted_date
        row_cells[3].text = "Issued for verification records"
        row_cells[4].text = "Basic Design"
        row_cells[5].text = "Quality Eng"
        row_cells[6].text = "Approvals Head"
        for cell in row_cells:
            cell.paragraphs[0].style.font.size = Pt(9)

        p_disc_lbl = doc.add_paragraph()
        p_disc_lbl.paragraph_format.space_before = Pt(40)
        r_disc_lbl = p_disc_lbl.add_run("DISCLAIMER")
        r_disc_lbl.font.bold = True
        r_disc_lbl.font.size = Pt(10)

        p_disc_txt = doc.add_paragraph()
        r_disc_txt = p_disc_txt.add_run(
            "The information and results of calculations included in this report are preliminary only. "
            "The results and calculations are to be checked and confirmed during award of contract and "
            "prior to construction of the vessel."
        )
        r_disc_txt.font.italic = True
        r_disc_txt.font.size = Pt(9.5)
        r_disc_txt.font.color.rgb = COLOR_CHARCOAL

        # ==========================================================
        # PAGE 3: TABLE OF CONTENTS
        # ==========================================================
        doc.add_page_break()
        append_running_page_header(doc)

        p_toc_head = doc.add_paragraph()
        r_toc_head = p_toc_head.add_run("Table of Contents")
        r_toc_head.font.name = 'Arial'
        r_toc_head.font.size = Pt(14)
        r_toc_head.font.bold = True
        r_toc_head.font.color.rgb = COLOR_DARK_TEXT

        table_toc = doc.add_table(rows=3, cols=3)
        table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_custom_table_borders(table_toc, HEX_BORDER_BLUE)
        table_toc.columns[0].width = Inches(1.0)
        table_toc.columns[1].width = Inches(5.0)
        table_toc.columns[2].width = Inches(1.0)

        toc_items = [
            ["1", "SECTION MODULUS ANALYSIS", "1-1"],
            ["1.1", "Input Dimensional Data Configuration", "1-1"],
            ["1.2", "Calculated Geometrical Results Matrix", "1-1"]
        ]

        for idx, item in enumerate(toc_items):
            row = table_toc.rows[idx]
            for col_idx, text in enumerate(item):
                row.cells[col_idx].text = text
                p = row.cells[col_idx].paragraphs[0]
                p.style.font.size = Pt(10)
                if idx == 0:
                    p.runs[0].font.bold = True

        # ==========================================================
        # PAGE 4: ENGINEERING DATA OUTPUT REPORT
        # ==========================================================
        doc.add_page_break()
        append_running_page_header(doc)

        p_report_title = doc.add_paragraph()
        r_rep_title = p_report_title.add_run("1 SECTION MODULUS ANALYSIS")
        r_rep_title.font.name = 'Arial'
        r_rep_title.font.size = Pt(15)
        r_rep_title.font.bold = True
        r_rep_title.font.color.rgb = COLOR_BLUE

        p_sub_section = doc.add_paragraph()
        r_sub_sec = p_sub_section.add_run("1.1 Input Properties Configuration Matrix")
        r_sub_sec.font.bold = True
        r_sub_sec.font.size = Pt(11)

        table_in_props = doc.add_table(rows=5, cols=2)
        set_custom_table_borders(table_in_props, HEX_BORDER_BLUE)
        
        input_dataset = [
            ["INPUTS Parameter Variables", "Engineered Configuration Value"],
            ["Selected Profile Type", str(profile)],
            ["Web Height (hw) / Bar Height (h)", f"{web_height} mm"],
            ["Web Thickness (tw) / Bar Thickness (t)", f"{web_thickness} mm"],
            ["Flange Width (bf)", "N/A" if profile == "Flat Bar" else f"{flange_width} mm"]
        ]

        for idx, row_pair in enumerate(input_dataset):
            row = table_in_props.rows[idx]
            row.cells[0].text = row_pair[0]
            row.cells[1].text = row_pair[1]
            if idx == 0:
                set_cell_shading(row.cells[0], HEX_LIGHT_BG)
                set_cell_shading(row.cells[1], HEX_LIGHT_BG)
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[1].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].style.font.size = Pt(10)
            row.cells[1].paragraphs[0].style.font.size = Pt(10)

        p_calc_lbl = doc.add_paragraph()
        p_calc_lbl.paragraph_format.space_before = Pt(20)
        r_calc_lbl = p_calc_lbl.add_run("1.2 Calculated Results Matrix")
        r_calc_lbl.font.bold = True
        r_calc_lbl.font.size = Pt(11)

        table_results = doc.add_table(rows=5, cols=2)
        set_custom_table_borders(table_results, HEX_BORDER_BLUE)

        results_dataset = [
            ["RESULTS Technical Nomenclature", "Computed Output Property Value"],
            ["Section Modulus (Bottom)", f"{z_bottom} cm³"],
            ["Total Cross-Sectional Area", f"{area} cm²"],
            ["Centroid Vertical Position (y-bar)", f"{centroid} mm"],
            ["Moment of Inertia (Ix)", f"{moment_of_inertia} cm⁴"]
        ]

        for idx, row_pair in enumerate(results_dataset):
            row = table_results.rows[idx]
            row.cells[0].text = row_pair[0]
            row.cells[1].text = row_pair[1]
            if idx == 0:
                set_cell_shading(row.cells[0], HEX_LIGHT_BG)
                set_cell_shading(row.cells[1], HEX_LIGHT_BG)
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[1].paragraphs[0].runs[0].font.bold = True
            elif idx == 1:
                row.cells[1].paragraphs[0].runs[0].font.bold = True
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = COLOR_BLUE
            row.cells[0].paragraphs[0].style.font.size = Pt(10)
            row.cells[1].paragraphs[0].style.font.size = Pt(10)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"CSL_Section_Properties_Report_{profile.replace(' ', '_')}.docx"
        )
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500