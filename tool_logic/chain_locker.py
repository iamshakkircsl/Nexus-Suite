"""
tool_logic/chain_locker.py
Chain Locker Sizing — All 5 Methods (Cleaned Production Release)
"""

import os
import copy
import tempfile
import numpy as np
import openpyxl
import formulas
from flask import Blueprint, request, jsonify, send_file
import threading
import io
import datetime

# ── WORD DOCUMENT FORMATTING IMPORTS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ── Blueprint Registration
chain_locker_bp = Blueprint('chain_locker', __name__)

TEMPLATE_PATH = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', 'excel_templates', 'chain_locker_template.xlsx'
))

BOOK = os.path.basename(TEMPLATE_PATH)

_MODEL_LOCK = threading.Lock()
_XL_MODEL   = formulas.ExcelModel().loads(TEMPLATE_PATH).finish()


# ════════════════════════════════════════════════
# CORE HELPERS
# ════════════════════════════════════════════════

def _ref(sheet: str, cell: str) -> str:
    return f"'[{BOOK}]{sheet.upper()}'!{cell}"

def _run(sheet: str, input_map: dict, result_cells: list) -> dict:
    with _MODEL_LOCK:
        inputs = {
            _ref(sheet, cell): np.array([[val]])
            for cell, val in input_map.items()
        }
        outputs = [_ref(sheet, c) for c in result_cells]
        solution = _XL_MODEL.calculate(inputs=inputs, outputs=outputs)
        results = {}
        for cell in result_cells:
            key = _ref(sheet, cell)
            try:
                raw = solution[key].value
                if hasattr(raw, '__iter__'):
                    raw = list(raw)[0]
                if hasattr(raw, '__iter__'):
                    raw = list(raw)[0]
                results[cell] = raw
            except (KeyError, IndexError):
                results[cell] = None
        return results

def _read_lookup(sheet_name: str, min_row: int, e_col: int, f_col: int) -> dict:
    wb = openpyxl.load_workbook(TEMPLATE_PATH, data_only=True)
    ws = wb[sheet_name]
    table = {}
    for row in ws.iter_rows(min_row=min_row, min_col=e_col, max_col=f_col, values_only=True):
        if row[0] is not None:
            table[float(row[0])] = float(row[1])
    return table

def _nearest_lookup(table: dict, value: float) -> float:
    nearest_key = min(table.keys(), key=lambda x: abs(x - value))
    return table[nearest_key]

def _fill_excel_export(sheet_names: list, input_sheets_cells: dict) -> str:
    wb_src = openpyxl.load_workbook(TEMPLATE_PATH)
    for sheet_name, cell_map in input_sheets_cells.items():
        ws = wb_src[sheet_name]
        for cell, val in cell_map.items():
            ws[cell] = val

    wb_out = openpyxl.Workbook()
    wb_out.remove(wb_out.active)

    for sheet_name in sheet_names:
        ws_src = wb_src[sheet_name]
        ws_out = wb_out.create_sheet(title=sheet_name)
        for row in ws_src.iter_rows():
            for cell in row:
                new_cell = ws_out.cell(row=cell.row, column=cell.column, value=cell.value)
                if cell.has_style:
                    new_cell.font          = copy.copy(cell.font)
                    new_cell.fill          = copy.copy(cell.fill)
                    new_cell.border        = copy.copy(cell.border)
                    new_cell.alignment     = copy.copy(cell.alignment)
                    new_cell.number_format = cell.number_format
        for col, dim in ws_src.column_dimensions.items():
            ws_out.column_dimensions[col].width = dim.width

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx', prefix='cl_export_')
    wb_out.save(tmp.name)
    tmp.close()
    return tmp.name

def _fmt(val, decimals=3, unit=''):
    if val is None:
        return '—'
    try:
        return f"{round(float(val), decimals)}{' ' + unit if unit else ''}"
    except (TypeError, ValueError):
        return str(val)

# ════════════════════════════════════════════════
# METHOD INPUT MAPPING FUNCTIONS
# ════════════════════════════════════════════════

def _ihc_input_map(d):
    return {
        'B3': float(d['chain_dia']),
        'B4': float(d['chain_length']),
        'B5': float(d['locker_dia']),
        'B6': d.get('taper', 'No'),
        'B7': d.get('mudbox', 'No'),
    }

def _cosco_round_input_map(d):
    return {
        'B3': float(d['chain_dia']),
        'B4': float(d['ratio']),
        'B5': float(d['chain_length']),
        'B6': d.get('taper', 'No'),
        'B7': d.get('mudbox', 'No'),
    }

def _cosco_rect_input_map(d):
    return {
        'B3': float(d['chain_dia']),
        'B4': float(d['chain_length']),
        'B5': float(d['locker_length']),
        'B6': float(d['locker_breadth']),
        'B7': float(d['k1']),
        'B8': d.get('taper', 'No'),
        'B9': d.get('mudbox', 'No'),
    }

def _ah_leng_run(d):
    chain_dia = float(d['chain_dia'])
    shape     = d.get('shape', 'Circle')
    lookup = _read_lookup('AH_Leng', min_row=3, e_col=5, f_col=6)
    vol_per_100m = _nearest_lookup(lookup, chain_dia)

    input_map = {
        'B2':  chain_dia,
        'B3':  float(d['chain_length']),
        'B4':  shape,
        'B5':  float(d.get('locker_dia', 0))     if shape == 'Circle'    else 0,
        'B6':  float(d.get('locker_length', 0))  if shape == 'Rectangle' else 0,
        'B7':  float(d.get('locker_breadth', 0)) if shape == 'Rectangle' else 0,
        'B8':  d.get('taper', 'No'),
        'B9':  d.get('mudbox', 'No'),
        'B12': vol_per_100m,
    }
    return _run('AH_Leng', input_map, ['B23', 'B24'])

def _ah_wt_run(d):
    chain_dia = float(d['chain_dia'])
    shape     = d.get('shape', 'Circle')
    lookup = _read_lookup('AH_Wt', min_row=3, e_col=5, f_col=6)
    wt_per_100m = _nearest_lookup(lookup, chain_dia)

    input_map = {
        'B2':  chain_dia,
        'B3':  float(d['chain_length']),
        'B4':  float(d.get('packing', 0.5)),
        'B5':  shape,
        'B6':  float(d.get('locker_dia', 0))     if shape == 'Circle'    else 0,
        'B7':  float(d.get('locker_length', 0))  if shape == 'Rectangle' else 0,
        'B8':  float(d.get('locker_breadth', 0)) if shape == 'Rectangle' else 0,
        'B9':  d.get('taper', 'No'),
        'B10': d.get('mudbox', 'No'),
        'B13': wt_per_100m,
    }
    return _run('AH_Wt', input_map, ['B25', 'B26'])

# ════════════════════════════════════════════════
# CORE METHOD ACTION ROUTES
# ════════════════════════════════════════════════

@chain_locker_bp.route('/api/ihc/calculate', methods=['POST'])
def ihc_calculate():
    d = request.get_json()
    try:
        r = _run('IHC_Calc', _ihc_input_map(d), ['B24', 'B25', 'B26'])
        return jsonify({'ok': True, 'results': {
            'volume':  r['B24'],
            'height':  r['B25'],
            'dh_text': r['B26'],
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/cosco-round/calculate', methods=['POST'])
def cosco_round_calculate():
    d = request.get_json()
    try:
        r = _run('COSCO_Round', _cosco_round_input_map(d), ['B26', 'B27'])
        return jsonify({'ok': True, 'results': {
            'volume': r['B26'],
            'height': r['B27'],
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/cosco-rect/calculate', methods=['POST'])
def cosco_rect_calculate():
    d = request.get_json()
    try:
        r = _run('COSCO_Rect', _cosco_rect_input_map(d), ['B24', 'B25', 'B26'])
        return jsonify({'ok': True, 'results': {
            'volume':  r['B24'],
            'height':  r['B25'],
            'bh_text': r['B26'],
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/ah-length/calculate', methods=['POST'])
def ah_leng_calculate():
    d = request.get_json()
    try:
        r = _ah_leng_run(d)
        return jsonify({'ok': True, 'results': {
            'volume': r['B23'],
            'height': r['B24'],
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/ah-weight/calculate', methods=['POST'])
def ah_wt_calculate():
    d = request.get_json()
    try:
        r = _ah_wt_run(d)
        return jsonify({'ok': True, 'results': {
            'volume': r['B25'],
            'height': r['B26'],
        }})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

# ════════════════════════════════════════════════
# EXCEL SPREADSHEET EXPORT WORKFLOWS
# ════════════════════════════════════════════════

@chain_locker_bp.route('/api/ihc/export-excel', methods=['POST'])
def ihc_export_excel():
    d = request.form.to_dict()
    try:
        path = _fill_excel_export(['IHC_Calc', 'IHC_Data'], {'IHC_Calc': _ihc_input_map(d)})
        return send_file(path, as_attachment=True, download_name='chain_locker_IHC.xlsx')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/cosco-round/export-excel', methods=['POST'])
def cosco_round_export_excel():
    d = request.form.to_dict()
    try:
        path = _fill_excel_export(['COSCO_Round'], {'COSCO_Round': _cosco_round_input_map(d)})
        return send_file(path, as_attachment=True, download_name='chain_locker_COSCO_Round.xlsx')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/cosco-rect/export-excel', methods=['POST'])
def cosco_rect_export_excel():
    d = request.form.to_dict()
    try:
        path = _fill_excel_export(['COSCO_Rect'], {'COSCO_Rect': _cosco_rect_input_map(d)})
        return send_file(path, as_attachment=True, download_name='chain_locker_COSCO_Rect.xlsx')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/ah-length/export-excel', methods=['POST'])
def ah_leng_export_excel():
    d = request.form.to_dict()
    try:
        shape = d.get('shape', 'Circle')
        chain_dia = float(d['chain_dia'])
        lookup = _read_lookup('AH_Leng', min_row=3, e_col=5, f_col=6)
        cell_map = {
            'B2':  chain_dia,
            'B3':  float(d['chain_length']),
            'B4':  shape,
            'B5':  float(d.get('locker_dia', 0))     if shape == 'Circle'    else 0,
            'B6':  float(d.get('locker_length', 0))  if shape == 'Rectangle' else 0,
            'B7':  float(d.get('locker_breadth', 0)) if shape == 'Rectangle' else 0,
            'B8':  d.get('taper', 'No'),
            'B9':  d.get('mudbox', 'No'),
            'B12': _nearest_lookup(lookup, chain_dia),
        }
        path = _fill_excel_export(['AH_Leng'], {'AH_Leng': cell_map})
        return send_file(path, as_attachment=True, download_name='chain_locker_AH_Length.xlsx')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/ah-weight/export-excel', methods=['POST'])
def ah_wt_export_excel():
    d = request.form.to_dict()
    try:
        shape = d.get('shape', 'Circle')
        chain_dia = float(d['chain_dia'])
        lookup = _read_lookup('AH_Wt', min_row=3, e_col=5, f_col=6)
        cell_map = {
            'B2':  chain_dia,
            'B3':  float(d['chain_length']),
            'B4':  float(d.get('packing', 0.5)),
            'B5':  shape,
            'B6':  float(d.get('locker_dia', 0))     if shape == 'Circle'    else 0,
            'B7':  float(d.get('locker_length', 0))  if shape == 'Rectangle' else 0,
            'B8':  float(d.get('locker_breadth', 0)) if shape == 'Rectangle' else 0,
            'B9':  d.get('taper', 'No'),
            'B10': d.get('mudbox', 'No'),
            'B13': _nearest_lookup(lookup, chain_dia),
        }
        path = _fill_excel_export(['AH_Wt'], {'AH_Wt': cell_map})
        return send_file(path, as_attachment=True, download_name='chain_locker_AH_Weight.xlsx')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

# ════════════════════════════════════════════════
# UNIFIED COCHIN SHIPYARD REPORTLAB PDF GENERATOR
# ════════════════════════════════════════════════

# ════════════════════════════════════════════════
# UNIFIED COCHIN SHIPYARD REPORTLAB PDF GENERATOR (UPDATED EXACT LAYOUT)
# ════════════════════════════════════════════════

# ════════════════════════════════════════════════
# UNIFIED COCHIN SHIPYARD REPORTLAB PDF GENERATOR (UPDATED EXACT LAYOUT)
# ════════════════════════════════════════════════

# ════════════════════════════════════════════════
# UNIFIED COCHIN SHIPYARD REPORTLAB PDF GENERATOR (UPDATED EXACT LAYOUT)
# ════════════════════════════════════════════════

# ════════════════════════════════════════════════
# UNIFIED COCHIN SHIPYARD REPORTLAB PDF GENERATOR (UPDATED EXACT LAYOUT)
# ════════════════════════════════════════════════

def _make_pdf(method_label: str, input_rows: list, result_rows: list) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', prefix='cl_report_')
    pdf_path = tmp.name
    tmp.close()

    # Using 1.9cm margins as specified in backend layout rules
    doc = SimpleDocTemplate(
        pdf_path, pagesize=A4,
        leftMargin=1.9*cm, rightMargin=1.9*cm,
        topMargin=1.9*cm, bottomMargin=1.9*cm
    )
    styles = getSampleStyleSheet()
    
    # Exact CSL Corporate Color Palette Matching
    COLOR_BLUE = colors.HexColor('#2b78c5')
    COLOR_LIGHT_BLUE = colors.HexColor('#7cb5ec')
    COLOR_BORDER = colors.HexColor('#5b9bd5')
    COLOR_DARK = colors.HexColor('#111111')
    COLOR_CHARCOAL = colors.HexColor('#333333')
    COLOR_MUTED = colors.HexColor('#6e6e6e')
    COLOR_BG = colors.HexColor('#f2f7fc')

    # Path to your uploaded corporate logo
    BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    LOGO_IMG_PATH = os.path.join(BASE_DIR, 'static', 'images', 'csl_logo.png')

    # ── Background Canvas Drawing for Page 1 (Cover Page)
    def draw_cover_background(canvas, doc):
        canvas.saveState()
        
        # 1. Left Vertical Decorative Bars (As seen in Screenshot 1)
        canvas.setFillColor(COLOR_BLUE)
        canvas.rect(0, 0, 0.4*cm, A4[1], fill=True, stroke=False)
        canvas.rect(0.6*cm, 0, 0.15*cm, A4[1], fill=True, stroke=False)
        
        # 2. Horizontal Blue Separator Bar underneath the Top Banner Logo
        canvas.setStrokeColor(COLOR_LIGHT_BLUE)
        canvas.setLineWidth(1.5)
        canvas.line(1.9*cm, A4[1] - 4.2*cm, A4[0] - 1.9*cm, A4[1] - 4.2*cm)
        
        # 3. Dynamic Logo Positioning on Cover Page Banner Header (MADE BIGGER)
        try:
            canvas.drawImage(LOGO_IMG_PATH, 1.9*cm + 0.2*cm, A4[1] - 3.8*cm, width=3.2*cm, height=3.2*cm, mask='auto')
        except Exception:
            # Fallback border box if logo image asset is temporarily missing during migration
            canvas.setStrokeColor(COLOR_BLUE)
            canvas.rect(1.9*cm + 0.2*cm, A4[1] - 3.8*cm, 3.2*cm, 3.2*cm, fill=False, stroke=True)

        canvas.restoreState()

    # ── Running Header Canvas Drawing for Pages 2, 3, 4 (As seen in Screenshot 2)
    def draw_running_header(canvas, doc):
        canvas.saveState()
        
        # Text Configurations matching Screenshot 2 Left Hand Side
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(COLOR_CHARCOAL)
        canvas.drawString(1.9*cm, A4[1] - 1.2*cm, "-Drawing Title-")
        canvas.drawString(1.9*cm, A4[1] - 1.6*cm, "-Drawing No-")
        canvas.drawString(1.9*cm, A4[1] - 2.0*cm, "-Rev. /Date-")
        
        # Text Configurations matching Screenshot 2 Right Hand Side (Shifted left to accommodate larger logo)
        canvas.drawRightString(A4[0] - 4.8*cm, A4[1] - 1.2*cm, "-Project No.-")
        canvas.drawRightString(A4[0] - 4.8*cm, A4[1] - 1.6*cm, "-Yard No.-")
        canvas.drawRightString(A4[0] - 4.8*cm, A4[1] - 2.0*cm, "-Vessel/Project Name-")
        
        # Running Top Right Header Logo Integration (MADE BIGGER)
        try:
            canvas.drawImage(LOGO_IMG_PATH, A4[0] - 4.2*cm, A4[1] - 2.4*cm, width=2.0*cm, height=2.0*cm, mask='auto')
        except Exception:
            pass
        
        # Running Header Bottom Styled Rule Line (Slightly Gradient / Tinted Shadow Effect)
        canvas.setStrokeColor(COLOR_LIGHT_BLUE)
        canvas.setLineWidth(1.5)
        canvas.line(1.9*cm, A4[1] - 2.7*cm, A4[0] - 1.9*cm, A4[1] - 2.7*cm)
        canvas.restoreState()

    story = []
    formatted_date = datetime.date.today().strftime('%d-%m-%Y')

    # ════════════════════════ PAGE 1: COVER PAGE FLOWABLES ════════════════════════
    # Spacer offset to account for background absolute logo canvas coordinates
    story.append(Spacer(1, 0.4*cm))
    
    # Header Enterprise Titles (Pushed right to clear the newly upscaled logo)
    story.append(Paragraph("<font size=15><b>COCHIN SHIPYARD LTD.</b></font>", 
                           ParagraphStyle('C1', fontName='Helvetica-Bold', leftIndent=3.8*cm, textColor=COLOR_BLUE, spaceAfter=2)))
    story.append(Paragraph("<font size=9.5><b>A GOVT. OF INDIA ENTERPRISE</b></font>", 
                           ParagraphStyle('C2', fontName='Helvetica-Bold', leftIndent=3.8*cm, textColor=COLOR_MUTED)))
    
    story.append(Spacer(1, 4.5*cm))
    
    # Updated Project Name Header block to exact requested casing/nomenclature
    story.append(Paragraph("<b>NEXUS: MARINE DESIGN SUITE</b>", ParagraphStyle('P1', fontName='Helvetica-Bold', fontSize=24, alignment=1, textColor=COLOR_DARK)))
    
    story.append(Spacer(1, 0.5*cm))
    t_line = Table([['']], colWidths=[A4[0]-3.8*cm], rowHeights=[1.5])
    t_line.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,-1), COLOR_BLUE)]))
    story.append(t_line)
    
    story.append(Spacer(1, 4.5*cm))
    story.append(Paragraph("<i>-Insert Vessel Picture-</i>", ParagraphStyle('P2', fontName='Helvetica-Oblique', fontSize=11, alignment=1, textColor=COLOR_MUTED)))
    
    story.append(Spacer(1, 4.5*cm))
    story.append(Paragraph("<b>DRAWING TITLE</b>", ParagraphStyle('P3', fontName='Helvetica-Bold', fontSize=18, alignment=1, textColor=COLOR_DARK)))
    story.append(Spacer(1, 0.6*cm))
    
    # Meta Configuration Matrix Layout block
    meta_data = [
        ["Drg. No.", ": CSL-NA-CL-001", "Department", ": Naval Architecture"],
        ["Project", f": {method_label}", "Document Type", ": Sizing Report"],
        ["Yard No.", ": BY-1002", "Owner", ": Technical Records Section"],
        ["Class", ": IRS / DNV GL Dual Class", "Pages", f": {formatted_date}"]
    ]
    t_cover = Table(meta_data, colWidths=[3.2*cm, 5.4*cm, 3.2*cm, 5.4*cm])
    t_cover.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('TEXTCOLOR', (0,0), (-1,-1), COLOR_DARK),
        ('FONTNAME', (0,0), (0,-1), 'Helvetica-Bold'),
        ('FONTNAME', (2,0), (2,-1), 'Helvetica-Bold'),
        ('LINEABOVE', (0,0), (-1,0), 1, COLOR_BORDER),
        ('LINEBELOW', (0,-1), (-1,-1), 1, COLOR_BORDER),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, COLOR_BG),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_cover)
    
    story.append(Spacer(1, 1.2*cm))
    story.append(Paragraph(
        "<i>This document/specification is the property of Cochin Shipyard Limited, and it must not be copied "
        "or the contents thereof or any information received in conjunction therewith must not be imparted/shared "
        "to any third party or utilized for any other purpose. The receipt of the document/specification implies "
        "that the conditions as mentioned herein are accepted.</i>",
        ParagraphStyle('Prop', fontName='Helvetica-Oblique', fontSize=8.5, textColor=COLOR_CHARCOAL, alignment=4)
    ))
    
    # ════════════════════════ PAGE 2: REVISION SHEET ════════════════════════
    story.append(PageBreak())
    story.append(Spacer(1, 1.5*cm))  # Clear space beneath the running header rule line
    story.append(Paragraph("<b>Revision Control Sheet</b>", ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=14, textColor=COLOR_DARK)))
    story.append(Spacer(1, 0.4*cm))
    
    rev_headers = [["S. No", "Revision", "Date", "Description", "Prepared", "Checked", "Approved"]]
    rev_data = [["1.", "0", formatted_date, "Issued for sizing confirmation verification records", "Naval Arch", "Quality Eng", "Approvals Head"]]
    t_rev = Table(rev_headers + rev_data, colWidths=[1.2*cm, 1.5*cm, 2.2*cm, 5.5*cm, 2.3*cm, 2.3*cm, 2.4*cm])
    t_rev.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 9),
        ('BACKGROUND', (0,0), (-1,0), COLOR_BG),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_rev)
    
    story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph("<b>DISCLAIMER</b>", ParagraphStyle('DiscL', fontName='Helvetica-Bold', fontSize=10, textColor=COLOR_DARK)))
    story.append(Spacer(1, 0.1*cm))
    story.append(Paragraph(
        "<i>The information and results of calculations included in this report are preliminary only. "
        "The results and calculations are to be checked and confirmed during award of contract and prior to construction of the vessel.</i>",
        ParagraphStyle('DiscT', fontName='Helvetica-Oblique', fontSize=9.5, textColor=COLOR_CHARCOAL, alignment=4)
    ))

    # ════════════════════════ PAGE 3: TABLE OF CONTENTS ════════════════════════
    story.append(PageBreak())
    story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph("<b>Table of Contents</b>", ParagraphStyle('TOC_H', fontName='Helvetica-Bold', fontSize=14, textColor=COLOR_DARK)))
    story.append(Spacer(1, 0.5*cm))
    
    toc_data = [
        ["1", "<b>CHAIN LOCKER STOWAGE CAPACITY SIZING ANALYSIS</b>", "1-1"],
        ["1.1", f"Input Parametric Boundary Fields ({method_label})", "1-1"],
        ["1.2", "Computed Volumetric Sizing Dimensions Results Matrix", "1-1"]
    ]
    t_toc = Table(toc_data, colWidths=[1.5*cm, 14.0*cm, 1.9*cm])
    t_toc.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('ALIGN', (2,0), (2,-1), 'RIGHT'),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, COLOR_BORDER),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t_toc)

    # ════════════════════════ PAGE 4: ENGINEERING DATA METRIC REPORT ════════════════════════
    story.append(PageBreak())
    story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph("<b>1 CHAIN LOCKER CAPACITY ANALYSIS</b>", ParagraphStyle('Sec1', fontName='Helvetica-Bold', fontSize=15, textColor=COLOR_BLUE)))
    story.append(Spacer(1, 0.4*cm))
    story.append(Paragraph("<b>1.1 Input Properties Configuration Matrix</b>", ParagraphStyle('Sub1', fontName='Helvetica-Bold', fontSize=11, textColor=COLOR_DARK)))
    story.append(Spacer(1, 0.2*cm))
    
    in_table_data = [["INPUTS Parameter Variables", "Engineered Configuration Value"]] + input_rows
    t_inputs = Table(in_table_data, colWidths=[10.5*cm, 6.9*cm])
    t_inputs.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('BACKGROUND', (0,0), (-1,0), COLOR_BG),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('LINEABOVE', (0,0), (-1,0), 1, COLOR_BORDER),
        ('LINEBELOW', (0,-1), (-1,-1), 1, COLOR_BORDER),
        ('LINEBELOW', (0,0), (-1,0), 1, COLOR_BORDER),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_inputs)
    
    story.append(Spacer(1, 0.6*cm))
    story.append(Paragraph("<b>1.2 Calculated Results Matrix</b>", ParagraphStyle('Sub2', fontName='Helvetica-Bold', fontSize=11, textColor=COLOR_DARK)))
    story.append(Spacer(1, 0.2*cm))
    
    res_table_data = [["RESULTS Technical Nomenclature", "Computed Output Sizing Value"]] + result_rows
    t_results = Table(res_table_data, colWidths=[10.5*cm, 6.9*cm])
    t_results.setStyle(TableStyle([
        ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('BACKGROUND', (0,0), (-1,0), COLOR_BG),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('TEXTCOLOR', (1,1), (1,1), COLOR_BLUE),
        ('FONTNAME', (1,1), (1,1), 'Helvetica-Bold'),
        ('LINEABOVE', (0,0), (-1,0), 1, COLOR_BORDER),
        ('LINEBELOW', (0,-1), (-1,-1), 1, COLOR_BORDER),
        ('LINEBELOW', (0,0), (-1,0), 1, COLOR_BORDER),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_results)

    doc.build(story, onFirstPage=draw_cover_background, onLaterPages=draw_running_header)
    return pdf_path
# ════════════════════════════════════════════════
# ROUTE ENDPOINTS FOR PDF EXPORTS
# ════════════════════════════════════════════════

@chain_locker_bp.route('/api/ihc/export-pdf', methods=['POST'])
def ihc_export_pdf():
    d = request.form.to_dict()
    try:
        r = _run('IHC_Calc', _ihc_input_map(d), ['B24', 'B25', 'B26'])
        pdf = _make_pdf('IHC Yard Standard',
            [['Chain Diameter (d)',     f"{d['chain_dia']} mm"],
             ['Total Chain Length (L)', f"{d['chain_length']} m"],
             ['Locker Diameter (D)',    f"{d['locker_dia']} m"],
             ['30° Taper Included',    d.get('taper', 'No')],
             ['Mudbox Included',       d.get('mudbox', 'No')]],
            [['Required Locker Volume', _fmt(r['B24'], unit='m³')],
             ['Required Locker Height', _fmt(r['B25'], unit='m')],
             ['D/H Ratio Acceptable?',  str(r['B26'])]])
        return send_file(pdf, as_attachment=True, download_name='chain_locker_IHC.pdf', mimetype='application/pdf')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/cosco-round/export-pdf', methods=['POST'])
def cosco_round_export_pdf():
    d = request.form.to_dict()
    try:
        r   = _run('COSCO_Round', _cosco_round_input_map(d), ['B26', 'B27'])
        dia = round(float(d['chain_dia']) * float(d['ratio']) / 1000, 3)
        pdf = _make_pdf('COSCO Round Locker',
            [['Chain Diameter (d)',     f"{d['chain_dia']} mm"],
             ['D/d Ratio',             f"{d['ratio']} (D = {dia} m)"],
             ['Total Chain Length (L)', f"{d['chain_length']} m"],
             ['30° Taper Included',    d.get('taper', 'No')],
             ['Mudbox Included',       d.get('mudbox', 'No')]],
            [['Required Locker Volume', _fmt(r['B26'], unit='m³')],
             ['Required Locker Height', _fmt(r['B27'], unit='m')]])
        return send_file(pdf, as_attachment=True, download_name='chain_locker_COSCO_Round.pdf', mimetype='application/pdf')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/cosco-rect/export-pdf', methods=['POST'])
def cosco_rect_export_pdf():
    d = request.form.to_dict()
    try:
        r   = _run('COSCO_Rect', _cosco_rect_input_map(d), ['B24', 'B25', 'B26'])
        pdf = _make_pdf('COSCO Rectangular Locker',
            [['Chain Diameter (d)',      f"{d['chain_dia']} mm"],
             ['Total Chain Length (L)',  f"{d['chain_length']} m"],
             ['Locker Length',           f"{d['locker_length']} m"],
             ['Locker Breadth',          f"{d['locker_breadth']} m"],
             ['COSCO Coefficient (k1)', d['k1']],
             ['30° Taper Included',     d.get('taper', 'No')],
             ['Mudbox Included',        d.get('mudbox', 'No')]],
            [['Required Locker Volume', _fmt(r['B24'], unit='m³')],
             ['Required Locker Height', _fmt(r['B25'], unit='m')],
             ['B/H Ratio Acceptable?',  str(r['B26'])]])
        return send_file(pdf, as_attachment=True, download_name='chain_locker_COSCO_Rect.pdf', mimetype='application/pdf')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/ah-length/export-pdf', methods=['POST'])
def ah_leng_export_pdf():
    d = request.form.to_dict()
    try:
        r     = _ah_leng_run(d)
        shape = d.get('shape', 'Circle')
        shape_rows = (
            [['Locker Diameter (D)', f"{d.get('locker_dia','—')} m"]]
            if shape == 'Circle' else
            [['Locker Length', f"{d.get('locker_length','—')} m"],
             ['Locker Breadth', f"{d.get('locker_breadth','—')} m"]]
        )
        pdf = _make_pdf('Anchor Handling — Length Based',
            [['Chain Diameter (d)',     f"{d['chain_dia']} mm"],
             ['Total Chain Length (L)', f"{d['chain_length']} m"],
             ['Locker Shape',           shape],
             *shape_rows,
             ['30° Taper Included',    d.get('taper', 'No')],
             ['Mudbox Included',       d.get('mudbox', 'No')]],
            [['Required Locker Volume', _fmt(r['B23'], unit='m³')],
             ['Required Locker Height', _fmt(r['B24'], unit='m')]])
        return send_file(pdf, as_attachment=True, download_name='chain_locker_AH_Length.pdf', mimetype='application/pdf')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

@chain_locker_bp.route('/api/ah-weight/export-pdf', methods=['POST'])
def ah_wt_export_pdf():
    d = request.form.to_dict()
    try:
        r     = _ah_wt_run(d)
        shape = d.get('shape', 'Circle')
        shape_rows = (
            [['Locker Diameter (D)', f"{d.get('locker_dia','—')} m"]]
            if shape == 'Circle' else
            [['Locker Length', f"{d.get('locker_length','—')} m"],
             ['Locker Breadth', f"{d.get('locker_breadth','—')} m"]]
        )
        packing = 'Tight (0.5 m³/t)' if float(d.get('packing', 0.5)) == 0.5 else 'Loose (0.7 m³/t)'
        pdf = _make_pdf('Anchor Handling — Weight Based',
            [['Chain Diameter (d)',     f"{d['chain_dia']} mm"],
             ['Total Chain Length (L)', f"{d['chain_length']} m"],
             ['Packing Nature',         packing],
             ['Locker Shape',           shape],
             *shape_rows,
             ['30° Taper Included',    d.get('taper', 'No')],
             ['Mudbox Included',       d.get('mudbox', 'No')]],
            [['Required Locker Volume', _fmt(r['B25'], unit='m³')],
             ['Required Locker Height', _fmt(r['B26'], unit='m')]])
        return send_file(pdf, as_attachment=True, download_name='chain_locker_AH_Weight.pdf', mimetype='application/pdf')
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 400

# ════════════════════════════════════════════════
# UNIFIED WORD ARCHITECTURE PIPELINE
# ════════════════════════════════════════════════

@chain_locker_bp.route('/api/<method>/export-word', methods=['POST'])
def export_word(method):
    try:
        d = request.get_json(silent=True) or request.form.to_dict()
        if not d:
            return jsonify({'ok': False, 'error': 'No input data received'}), 400

        if method == 'ihc':
            method_label = "IHC Yard Standard"
            r = _run('IHC_Calc', _ihc_input_map(d), ['B24', 'B25', 'B26'])
            inputs_dataset = [
                ["Chain Diameter (d)", f"{d.get('chain_dia')} mm"],
                ["Total Chain Length (L)", f"{d.get('chain_length')} m"],
                ["Locker Diameter (D)", f"{d.get('locker_dia')} m"],
                ["30° Taper Included", str(d.get('taper', 'No'))],
                ["Mudbox Included", str(d.get('mudbox', 'No'))]
            ]
            results_dataset = [
                ["Required Locker Volume", _fmt(r.get('B24'), unit='m³')],
                ["Required Locker Height", _fmt(r.get('B25'), unit='m')],
                ["D/H Ratio Acceptable?", str(r.get('B26'))]
            ]
        elif method == 'cosco-round':
            method_label = "COSCO Round Locker"
            r = _run('COSCO_Round', _cosco_round_input_map(d), ['B26', 'B27'])
            dia_calc = round(float(d.get('chain_dia', 0)) * float(d.get('ratio', 0)) / 1000, 3)
            inputs_dataset = [
                ["Chain Diameter (d)", f"{d.get('chain_dia')} mm"],
                ["D/d Ratio", f"{d.get('ratio')} (D = {dia_calc} m)"],
                ["Total Chain Length (L)", f"{d.get('chain_length')} m"],
                ["30° Taper Included", str(d.get('taper', 'No'))],
                ["Mudbox Included", str(d.get('mudbox', 'No'))]
            ]
            results_dataset = [
                ["Required Locker Volume", _fmt(r.get('B26'), unit='m³')],
                ["Required Locker Height", _fmt(r.get('B27'), unit='m')]
            ]
        elif method == 'cosco-rect':
            method_label = "COSCO Rectangular Locker"
            r = _run('COSCO_Rect', _cosco_rect_input_map(d), ['B24', 'B25', 'B26'])
            inputs_dataset = [
                ["Chain Diameter (d)", f"{d.get('chain_dia')} mm"],
                ["Total Chain Length (L)", f"{d.get('chain_length')} m"],
                ["Locker Length", f"{d.get('locker_length')} m"],
                ["Locker Breadth", f"{d.get('locker_breadth')} m"],
                ["COSCO Coefficient (k1)", str(d.get('k1'))],
                ["30° Taper Included", str(d.get('taper', 'No'))],
                ["Mudbox Included", str(d.get('mudbox', 'No'))]
            ]
            results_dataset = [
                ["Required Locker Volume", _fmt(r.get('B24'), unit='m³')],
                ["Required Locker Height", _fmt(r.get('B25'), unit='m')],
                ["B/H Ratio Acceptable?", str(r.get('B26'))]
            ]
        elif method == 'ah-length':
            method_label = "Anchor Handling (Length Based)"
            r = _ah_leng_run(d)
            inputs_dataset = [
                ["Chain Diameter (d)", f"{d.get('chain_dia')} mm"],
                ["Total Chain Length (L)", f"{d.get('chain_length')} m"],
                ["Locker Shape", str(d.get('shape'))],
                ["Locker Diameter (D)", f"{d.get('locker_dia')} m" if d.get('shape') == 'Circle' else "N/A"],
                ["Locker Dimension (L x B)", f"{d.get('locker_length')}m x {d.get('locker_breadth')}m" if d.get('shape') == 'Rectangle' else "N/A"],
                ["30° Taper Included", str(d.get('taper', 'No'))],
                ["Mudbox Included", str(d.get('mudbox', 'No'))]
            ]
            results_dataset = [
                ["Required Locker Volume", _fmt(r.get('B23'), unit='m³')],
                ["Required Locker Height", _fmt(r.get('B24'), unit='m')]
            ]
        elif method == 'ah-weight':
            method_label = "Anchor Handling (Weight Based)"
            r = _ah_wt_run(d)
            packing_lbl = 'Tight (0.5 m³/t)' if float(d.get('packing', 0.5)) == 0.5 else 'Loose (0.7 m³/t)'
            inputs_dataset = [
                ["Chain Diameter (d)", f"{d.get('chain_dia')} mm"],
                ["Total Chain Length (L)", f"{d.get('chain_length')} m"],
                ["Packing Nature", packing_lbl],
                ["Locker Shape", str(d.get('shape'))],
                ["Locker Diameter (D)", f"{d.get('locker_dia')} m" if d.get('shape') == 'Circle' else "N/A"],
                ["Locker Dimension (L x B)", f"{d.get('locker_length')}m x {d.get('locker_breadth')}m" if d.get('shape') == 'Rectangle' else "N/A"],
                ["30° Taper Included", str(d.get('taper', 'No'))],
                ["Mudbox Included", str(d.get('mudbox', 'No'))]
            ]
            results_dataset = [
                ["Required Locker Volume", _fmt(r.get('B25'), unit='m³')],
                ["Required Locker Height", _fmt(r.get('B26'), unit='m')]
            ]
        else:
            return jsonify({'ok': False, 'error': f"Unknown frame: {method}"}), 400

        formatted_date = datetime.date.today().strftime('%d-%m-%Y')
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        COLOR_BLUE = RGBColor(43, 120, 197)        
        COLOR_DARK_TEXT = RGBColor(17, 17, 17)     
        COLOR_MUTED_TEXT = RGBColor(110, 110, 110) 
        HEX_BORDER_BLUE = "5B9BD5"                 
        HEX_LIGHT_BG = "F2F7FC"                    
        COLOR_CHARCOAL = RGBColor(51, 51, 51)

        def set_cell_shading(cell, fill_hex):
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)

        def set_custom_table_borders(table, border_hex):
            tblPr = table._tbl.tblPr
            borders_elm = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_hex}"/>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_hex}"/>'
                f'  <w:left w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
                f'  <w:insideV w:val="none"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders_elm)

        def append_running_page_header(document):
            tbl = document.add_table(rows=1, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = False
            tbl.columns[0].width = Inches(3.5)
            tbl.columns[1].width = Inches(3.5)

            cell_l = tbl.rows[0].cells[0]
            p_l = cell_l.paragraphs[0]
            p_l.paragraph_format.space_after = Pt(2)
            r_l = p_l.add_run("-Drawing Title-\n-Drawing No-\n-Rev. /Date-")
            r_l.font.size = Pt(8.5)
            r_l.font.color.rgb = COLOR_MUTED_TEXT

            cell_r = tbl.rows[0].cells[1]
            p_r = cell_r.paragraphs[0]
            p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_r.paragraph_format.space_after = Pt(2)
            r_r = p_r.add_run("-Project No.-\n-Yard No.-\n-Vessel/Project Name-  ")
            r_r.font.size = Pt(8.5)
            r_r.font.color.rgb = COLOR_MUTED_TEXT
            
            try:
                BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
                word_logo_path = os.path.join(BASE_DIR, 'static', 'images', 'csl_logo.png')
                p_r.add_run().add_picture(word_logo_path, width=Inches(0.45))
            except Exception:
                pass

            p_line = document.add_paragraph()
            p_line.paragraph_format.space_before = Pt(0)
            p_line.paragraph_format.space_after = Pt(18)
            p_line_run = p_line.add_run("━" * 65)
            p_line_run.font.color.rgb = COLOR_BLUE
            p_line_run.font.size = Pt(8)

        # Cover Page
        p_logo = doc.add_paragraph()
        try:
            BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
            word_logo_path = os.path.join(BASE_DIR, 'static', 'images', 'csl_logo.png')
            p_logo.add_run().add_picture(word_logo_path, width=Inches(0.95))
        except Exception:
            pass

        p_corp = doc.add_paragraph()
        r_corp = p_corp.add_run("COCHIN SHIPYARD LTD.\nA GOVT. OF INDIA ENTERPRISE")
        r_corp.font.name = 'Arial'
        r_corp.font.size = Pt(13)
        r_corp.font.bold = True
        r_corp.font.color.rgb = COLOR_BLUE
        
        p_line_top = doc.add_paragraph()
        r_line_top = p_line_top.add_run("━" * 65)
        r_line_top.font.color.rgb = COLOR_BLUE

        p_proj_title = doc.add_paragraph()
        p_proj_title.paragraph_format.space_before = Pt(45)
        p_proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_proj = p_proj_title.add_run("NEXUS: MARINE DESIGN SUITE")
        r_proj.font.name = 'Arial'
        r_proj.font.size = Pt(22)
        r_proj.font.bold = True
        r_proj.font.color.rgb = COLOR_DARK_TEXT

        p_vessel = doc.add_paragraph()
        p_vessel.paragraph_format.space_before = Pt(75)
        p_vessel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ves = p_vessel.add_run("-Insert Vessel Picture-")
        r_ves.font.name = 'Arial'
        r_ves.font.size = Pt(11)
        r_ves.font.italic = True
        r_ves.font.color.rgb = COLOR_MUTED_TEXT

        p_drg_title = doc.add_paragraph()
        p_drg_title.paragraph_format.space_before = Pt(75)
        p_drg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_drg = p_drg_title.add_run("CHAIN LOCKER SIZING REPORT")
        r_drg.font.name = 'Arial'
        r_drg.font.size = Pt(16)
        r_drg.font.bold = True
        r_drg.font.color.rgb = COLOR_DARK_TEXT

        table_cover = doc.add_table(rows=4, cols=4)
        table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_custom_table_borders(table_cover, HEX_BORDER_BLUE)
        table_cover.autofit = False
        table_cover.columns[0].width = Inches(1.5)
        table_cover.columns[1].width = Inches(2.0)
        table_cover.columns[2].width = Inches(1.5)
        table_cover.columns[3].width = Inches(2.0)
        
        meta_matrix = [
            ["Drg. No.", ": CSL-NA-CL-001", "Department", ": Naval Architecture"],
            ["Project", f": {method_label}", "Document Type", ": Sizing Report"],
            ["Yard No.", ": BY-1002", "Owner", ": Technical Records Section"],
            ["Class", ": IRS / DNV GL Dual Class", "Pages", f": {formatted_date}"]
        ]
        
        for idx, row_vals in enumerate(meta_matrix):
            row = table_cover.rows[idx]
            for c_idx, val_text in enumerate(row_vals):
                cell = row.cells[c_idx]
                cell.text = val_text
                p = cell.paragraphs[0]
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(9.5)
                if c_idx % 2 == 0:
                    p.runs[0].font.bold = True

        p_disc_notice = doc.add_paragraph()
        p_disc_notice.paragraph_format.space_before = Pt(35)
        p_disc_notice.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_notice = p_disc_notice.add_run(
            "This document/specification is the property of Cochin Shipyard Limited, and it must not be copied "
            "or the contents thereof or any information received in conjunction therewith must not be imparted/shared "
            "to any third party or utilized for any other purpose. The receipt of the document/specification implies "
            "that the conditions as mentioned herein are accepted."
        )
        r_notice.font.size = Pt(8.5)
        r_notice.font.italic = True
        r_notice.font.color.rgb = COLOR_CHARCOAL

        # Page 2: Revision Sheet
        doc.add_page_break()
        append_running_page_header(doc)

        p_rev_head = doc.add_paragraph()
        r_rev_head = p_rev_head.add_run("Revision Control Sheet")
        r_rev_head.font.name = 'Arial'
        r_rev_head.font.size = Pt(14)
        r_rev_head.font.bold = True
        r_rev_head.font.color.rgb = COLOR_DARK_TEXT

        table_rev = doc.add_table(rows=2, cols=7)
        table_rev.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_custom_table_borders(table_rev, "000000")
        
        rev_headers = ["S. No", "Revision", "Date", "Description", "Prepared", "Checked", "Approved"]
        for idx, header_text in enumerate(rev_headers):
            cell = table_rev.rows[0].cells[idx]
            cell.text = header_text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_shading(cell, HEX_LIGHT_BG)

        row_cells = table_rev.rows[1].cells
        row_cells[0].text = "1."
        row_cells[1].text = "0"
        row_cells[2].text = formatted_date
        row_cells[3].text = "Issued for sizing confirmation verification records"
        row_cells[4].text = "Naval Arch"
        row_cells[5].text = "Quality Eng"
        row_cells[6].text = "Approvals Head"
        for cell in row_cells:
            cell.paragraphs[0].style.font.size = Pt(9)

        p_disc_lbl = doc.add_paragraph()
        p_disc_lbl.paragraph_format.space_before = Pt(40)
        r_disc_lbl = p_disc_lbl.add_run("DISCLAIMER")
        r_disc_lbl.font.bold = True
        r_disc_lbl.font.size = Pt(10)

        p_disc_txt = doc.add_paragraph()
        r_disc_txt = p_disc_txt.add_run(
            "The information and results of calculations included in this report are preliminary only. "
            "The results and calculations are to be checked and confirmed during award of contract and "
            "prior to construction of the vessel."
        )
        r_disc_txt.font.italic = True
        r_disc_txt.font.size = Pt(9.5)
        r_disc_txt.font.color.rgb = COLOR_CHARCOAL

        # Page 3: TOC
        doc.add_page_break()
        append_running_page_header(doc)

        p_toc_head = doc.add_paragraph()
        r_toc_head = p_toc_head.add_run("Table of Contents")
        r_toc_head.font.name = 'Arial'
        r_toc_head.font.size = Pt(14)
        r_toc_head.font.bold = True
        r_toc_head.font.color.rgb = COLOR_DARK_TEXT

        table_toc = doc.add_table(rows=3, cols=3)
        table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_custom_table_borders(table_toc, HEX_BORDER_BLUE)
        table_toc.columns[0].width = Inches(1.0)
        table_toc.columns[1].width = Inches(5.0)
        table_toc.columns[2].width = Inches(1.0)

        toc_items = [
            ["1", "CHAIN LOCKER STOWAGE CAPACITY SIZING ANALYSIS", "1-1"],
            ["1.1", f"Input Parametric Boundary Fields ({method_label})", "1-1"],
            ["1.2", "Computed Volumetric Sizing Dimensions Results Matrix", "1-1"]
        ]

        for idx, item in enumerate(toc_items):
            row = table_toc.rows[idx]
            for col_idx, text in enumerate(item):
                row.cells[col_idx].text = text
                p = row.cells[col_idx].paragraphs[0]
                p.style.font.size = Pt(10)
                if idx == 0:
                    p.runs[0].font.bold = True

        # Page 4: Summary Output Matrix
        doc.add_page_break()
        append_running_page_header(doc)

        p_report_title = doc.add_paragraph()
        r_rep_title = p_report_title.add_run("1 CHAIN LOCKER CAPACITY ANALYSIS")
        r_rep_title.font.name = 'Arial'
        r_rep_title.font.size = Pt(15)
        r_rep_title.font.bold = True
        r_rep_title.font.color.rgb = COLOR_BLUE

        p_sub_section = doc.add_paragraph()
        r_sub_sec = p_sub_section.add_run("1.1 Input Properties Configuration Matrix")
        r_sub_sec.font.bold = True
        r_sub_sec.font.size = Pt(11)

        table_in_props = doc.add_table(rows=len(inputs_dataset) + 1, cols=2)
        set_custom_table_borders(table_in_props, HEX_BORDER_BLUE)
        
        table_in_props.rows[0].cells[0].text = "INPUTS Parameter Variables"
        table_in_props.rows[0].cells[1].text = "Engineered Configuration Value"
        set_cell_shading(table_in_props.rows[0].cells[0], HEX_LIGHT_BG)
        set_cell_shading(table_in_props.rows[0].cells[1], HEX_LIGHT_BG)
        table_in_props.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table_in_props.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

        for idx, row_pair in enumerate(inputs_dataset):
            row = table_in_props.rows[idx + 1]
            row.cells[0].text = str(row_pair[0])
            row.cells[1].text = str(row_pair[1])
            row.cells[0].paragraphs[0].style.font.size = Pt(10)
            row.cells[1].paragraphs[0].style.font.size = Pt(10)

        p_calc_lbl = doc.add_paragraph()
        p_calc_lbl.paragraph_format.space_before = Pt(20)
        r_calc_lbl = p_calc_lbl.add_run("1.2 Calculated Results Matrix")
        r_calc_lbl.font.bold = True
        r_calc_lbl.font.size = Pt(11)

        table_results = doc.add_table(rows=len(results_dataset) + 1, cols=2)
        set_custom_table_borders(table_results, HEX_BORDER_BLUE)

        table_results.rows[0].cells[0].text = "RESULTS Technical Nomenclature"
        table_results.rows[0].cells[1].text = "Computed Output Sizing Value"
        set_cell_shading(table_results.rows[0].cells[0], HEX_LIGHT_BG)
        set_cell_shading(table_results.rows[0].cells[1], HEX_LIGHT_BG)
        table_results.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table_results.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

        for idx, row_pair in enumerate(results_dataset):
            row = table_results.rows[idx + 1]
            row.cells[0].text = str(row_pair[0])
            row.cells[1].text = str(row_pair[1])
            if idx == 0:
                row.cells[1].paragraphs[0].runs[0].font.bold = True
                row.cells[1].paragraphs[0].runs[0].font.color.rgb = COLOR_BLUE
            row.cells[0].paragraphs[0].style.font.size = Pt(10)
            row.cells[1].paragraphs[0].style.font.size = Pt(10)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        safe_name = method_label.replace(' ', '_').replace('(', '').replace(')', '')
        return send_file(
            file_stream,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"CSL_Chain_Locker_Report_{safe_name}.docx"
        )
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500